#!/usr/bin/env python3
"""
VASSWEB / VASS & CO. — Complete Document Rebrand Script
Brand: Black (#0a0a0a) + Gold (#d4a843) + White + Inter font
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, copy

# ─── BRAND CONSTANTS ───
BLACK = RGBColor(0x0a, 0x0a, 0x0a)
DARK = RGBColor(0x11, 0x11, 0x11)
GOLD = RGBColor(0xd4, 0xa8, 0x43)
GOLD_DARK = RGBColor(0xc9, 0x98, 0x1a)
WHITE = RGBColor(0xff, 0xff, 0xff)
GRAY = RGBColor(0x99, 0x99, 0x99)
LIGHT_GRAY = RGBColor(0xcc, 0xcc, 0xcc)
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)

FONT_NAME = "Inter"
FONT_HEADING = "Inter"

BASE = "/Users/vass/Library/Mobile Documents/com~apple~CloudDocs/vassweb-sk"
LOGO_H = f"{BASE}/public/images/logo-horizontal.png"
LOGO_ICON = f"{BASE}/public/images/logo-icon.png"
DOCS_DIR = f"{BASE}/vassweb-dokumenty"
DRIVE = "/Users/vass/Library/Mobile Documents/com~apple~CloudDocs/Vass&Co. drive"

COMPANY_INFO = {
    "name": "Vass & Co. s.r.o.",
    "brand": "VASSWEB",
    "email": "richard.vass@vassco.sk",
    "web": "vassweb.sk",
    "phone": "+421 918 668 728",
    "ico": "56921021",
    "dic": "2122501524",
    "ic_dph": "SK2122501524",
    "iban": "SK11 0900 0000 0052 3252 7162",
    "address": "Bratislava, Slovenská republika",
    "konatel": "Richard Vass",
}


# ─── HELPER FUNCTIONS ───

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            color, sz = val
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, font_size=10, bold=False, color=BODY_TEXT, alignment=None, space_before=0, space_after=6, font_name=FONT_NAME):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def add_gold_heading(doc, text, level=1):
    """Add a gold heading."""
    sizes = {1: 18, 2: 14, 3: 12}
    p = add_paragraph(doc, text, font_size=sizes.get(level, 12), bold=True, color=GOLD_DARK, space_before=16, space_after=8)
    return p


def add_header_block(doc, title_text=None):
    """Add branded header with logo and company info."""
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    header.is_linked_to_previous = False

    # Clear existing header
    for p in header.paragraphs:
        p.clear()

    # Header table: logo left, info right
    htable = header.add_table(rows=1, cols=2, width=Inches(6.5))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table borders
    for row in htable.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:left {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:bottom {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:right {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>')
            tcPr.append(tcBorders)

    # Logo cell
    logo_cell = htable.cell(0, 0)
    logo_cell.width = Inches(3)
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_H):
        run = lp.add_run()
        run.add_picture(LOGO_H, height=Cm(1.2))

    # Info cell
    info_cell = htable.cell(0, 1)
    info_cell.width = Inches(3.5)
    ip = info_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, line in enumerate([
        COMPANY_INFO["name"],
        f'{COMPANY_INFO["email"]}  |  {COMPANY_INFO["web"]}',
        COMPANY_INFO["phone"],
    ]):
        if i > 0:
            ip.add_run("\n")
        run = ip.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = Pt(7.5)
        run.font.color.rgb = GRAY

    # Gold line under header
    border_p = header.add_paragraph()
    border_p.paragraph_format.space_before = Pt(6)
    border_p.paragraph_format.space_after = Pt(0)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D4A843"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_footer_block(doc):
    """Add branded footer with gold line and company details."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    # Gold line
    border_p = footer.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)
    border_p.paragraph_format.space_after = Pt(4)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="1" w:color="D4A843"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Footer text
    info_text = (
        f'{COMPANY_INFO["name"]}  |  IČO: {COMPANY_INFO["ico"]}  |  '
        f'DIČ: {COMPANY_INFO["dic"]}  |  IČ DPH: {COMPANY_INFO["ic_dph"]}  |  '
        f'{COMPANY_INFO["address"]}'
    )
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(info_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY


def make_branded_table(doc, headers, rows, col_widths=None):
    """Create a branded table with gold header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "0A0A0A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = FONT_NAME
        run.font.size = Pt(8.5)
        run.bold = True
        run.font.color.rgb = GOLD

    # Style data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            bg = "FFFFFF" if ri % 2 == 0 else "F8F8F6"
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = BODY_TEXT

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    # Remove all borders, add subtle bottom borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:left {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:bottom {nsdecls("w")} w:val="single" w:sz="2" w:space="0" w:color="E8E0D0"/>'
                f'<w:right {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    return table


def add_signature_block(doc, left_name="Richard Vass", left_role="konateľ, Vass & Co. s.r.o.", right_name="[Meno a priezvisko]", right_role="Objednávateľ"):
    """Add signature block."""
    add_paragraph(doc, "", space_before=24)
    add_paragraph(doc, f"V Bratislave, dňa ___________________________", font_size=10, color=BODY_TEXT, space_after=20)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:left {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:bottom {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:right {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>')
            tcPr.append(tcBorders)

    for ci, (name, role) in enumerate([(left_name, left_role), (right_name, right_role)]):
        cell0 = table.cell(0, ci)
        p = cell0.paragraphs[0]
        run = p.add_run("______________________________")
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = LIGHT_GRAY

        cell1 = table.cell(1, ci)
        p = cell1.paragraphs[0]
        run = p.add_run(name)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = BODY_TEXT

        cell2 = table.cell(2, ci)
        p = cell2.paragraphs[0]
        run = p.add_run(role)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def add_parties_table(doc):
    """Add supplier/client info table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>')
        tcPr.append(tcBorders)

    # Left - supplier
    left = table.cell(0, 0)
    set_cell_shading(left, "F8F8F6")
    p = left.paragraphs[0]
    for i, line in enumerate(["DODÁVATEĽ:", "VASSWEB", "Vass & Co. s.r.o.", f"IČO: {COMPANY_INFO['ico']}", f"DIČ: {COMPANY_INFO['dic']}", f"IČ DPH: {COMPANY_INFO['ic_dph']}", f"IBAN: {COMPANY_INFO['iban']}", f"{COMPANY_INFO['email']}"]):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = BODY_TEXT
        if i == 0:
            run.bold = True
            run.font.color.rgb = GOLD_DARK
        elif i == 1:
            run.bold = True

    # Right - client
    right = table.cell(0, 1)
    set_cell_shading(right, "F8F8F6")
    p = right.paragraphs[0]
    for i, line in enumerate(["OBJEDNÁVATEĽ:", "[Meno a priezvisko / Spoločnosť]", "[Ulica a číslo]", "[PSČ Mesto]", "IČO: [IČO]", "DIČ: [DIČ]", "E-mail: [e-mail]"]):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = BODY_TEXT
        if i == 0:
            run.bold = True
            run.font.color.rgb = GOLD_DARK

    return table


def add_bullet_list(doc, items, color=BODY_TEXT):
    """Add styled bullet list."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        # Gold bullet
        run = p.add_run("●  ")
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = GOLD
        # Text
        run = p.add_run(item)
        run.font.name = FONT_NAME
        run.font.size = Pt(9.5)
        run.font.color.rgb = color


# ═══════════════════════════════════════════════════════════
# DOCUMENT GENERATORS
# ═══════════════════════════════════════════════════════════

def create_faktura():
    """Generate branded invoice template."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    # Title
    add_paragraph(doc, "FAKTÚRA", font_size=22, bold=True, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_paragraph(doc, "č. [ROK]-[ČÍSLO]", font_size=11, color=GRAY, space_after=16)

    # Parties
    add_parties_table(doc)
    add_paragraph(doc, "", space_after=8)

    # Invoice details
    make_branded_table(doc,
        headers=["Dátum vystavenia", "Dátum splatnosti", "Forma úhrady", "Variabilný symbol"],
        rows=[["[DD.MM.RRRR]", "[DD.MM.RRRR]", "Bankový prevod", "[číslo]"]],
    )
    add_paragraph(doc, "", space_after=8)

    # Items table
    make_branded_table(doc,
        headers=["POPIS", "MNOŽSTVO", "JEDN. CENA", "SPOLU"],
        rows=[
            ["[Popis služby]", "1", "[cena] €", "[cena] €"],
            ["[Popis služby]", "1", "[cena] €", "[cena] €"],
        ],
    )
    add_paragraph(doc, "", space_after=8)

    # Totals
    make_branded_table(doc,
        headers=["", ""],
        rows=[
            ["Základ DPH (20%):", "[suma] €"],
            ["DPH 20%:", "[suma] €"],
            ["CELKOM K ÚHRADE:", "[suma] €"],
        ],
    )

    doc.save(f"{DOCS_DIR}/Vassweb-faktura-sablona.docx")
    print("✓ Faktúra")


def create_cenova_ponuka():
    """Generate branded quote template."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "CENOVÁ PONUKA", font_size=22, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, "č. [ROK]-[ČÍSLO]", font_size=11, color=GRAY, space_after=16)

    # Details
    for label, value in [("Klient:", "[Meno klienta / Spoločnosť]"), ("Projekt:", "[Názov projektu]"), ("Dátum:", "[DD.MM.RRRR]"), ("Platnosť ponuky:", "30 dní od dátumu vystavenia")]:
        p = doc.add_paragraph()
        run = p.add_run(label + "  ")
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = BODY_TEXT
        run = p.add_run(value)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT

    add_paragraph(doc, "", space_after=8)
    add_gold_heading(doc, "Rozsah prác", level=2)

    make_branded_table(doc,
        headers=["POLOŽKA", "POPIS", "CENA (bez DPH)"],
        rows=[
            ["1.", "[Webová stránka / Aplikácia]", "[cena] €"],
            ["2.", "[Dizajn / UI/UX]", "[cena] €"],
            ["3.", "[SEO a optimalizácia]", "[cena] €"],
            ["4.", "[Hosting / Správa]", "[cena] €"],
        ],
    )

    add_paragraph(doc, "", space_after=8)
    p = doc.add_paragraph()
    run = p.add_run("Celková cena (bez DPH):  ")
    run.font.name = FONT_NAME; run.font.size = Pt(11); run.bold = True; run.font.color.rgb = BODY_TEXT
    run = p.add_run("[CELKOVÁ CENA] €")
    run.font.name = FONT_NAME; run.font.size = Pt(11); run.bold = True; run.font.color.rgb = GOLD_DARK

    p = doc.add_paragraph()
    run = p.add_run("DPH 20%:  ")
    run.font.name = FONT_NAME; run.font.size = Pt(10); run.font.color.rgb = BODY_TEXT
    run = p.add_run("[DPH] €")
    run.font.name = FONT_NAME; run.font.size = Pt(10); run.font.color.rgb = BODY_TEXT

    add_paragraph(doc, "", space_after=8)
    add_gold_heading(doc, "Platobné podmienky", level=2)
    add_bullet_list(doc, [
        "50% záloha pred zahájením prác",
        "50% doplatku po odovzdaní diela",
        f"IBAN: {COMPANY_INFO['iban']}",
    ])

    add_signature_block(doc, right_role="Objednávateľ")
    doc.save(f"{DOCS_DIR}/Vassweb-cenova-ponuka-sablona.docx")
    print("✓ Cenová ponuka")


def create_zmluva():
    """Generate branded contract template."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "ZMLUVA O DIELO", font_size=22, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, "uzatvorená podľa § 536 a nasl. Obchodného zákonníka SR", font_size=10, color=GRAY, space_after=16)

    add_parties_table(doc)
    add_paragraph(doc, "sa dohodli na uzatvorení tejto Zmluvy o dielo (ďalej len \u201eZmluva\u201c):", font_size=10, color=BODY_TEXT, space_before=12, space_after=12)

    articles = [
        ("Článok I – Predmet zmluvy", [
            "1. Zhotoviteľ sa zaväzuje vytvoriť pre Objednávateľa dielo: [popis diela], v rozsahu a podľa špecifikácie dohodnutej v prílohe tejto Zmluvy.",
            "2. Dielom sa rozumie výsledok tvorivej duševnej činnosti – webová stránka / aplikácia / iné digitálne dielo.",
        ]),
        ("Článok II – Cena a platobné podmienky", [
            "1. Celková cena za dielo je [suma] EUR bez DPH (resp. [suma] EUR s DPH 20%).",
            "2. Platba prebieha v dvoch splátkach: 50% záloha pred zahájením prác, 50% doplatku po odovzdaní diela.",
            f"3. Platba bezhotovostným prevodom na IBAN: {COMPANY_INFO['iban']}.",
        ]),
        ("Článok III – Termín plnenia", [
            "1. Zhotoviteľ sa zaväzuje odovzdať dielo do [počet] pracovných dní od uhradenia zálohy.",
            "2. V prípade omeškania Objednávateľa s poskytnutím podkladov sa termín primerane predlžuje.",
        ]),
        ("Článok IV – Autorské práva", [
            "1. Po úplnej úhrade ceny diela prechádzajú majetkové práva k dielu na Objednávateľa.",
            "2. Zhotoviteľ si vyhradzuje právo uviesť dielo vo svojom portfóliu s uvedením klienta.",
        ]),
        ("Článok V – Záručná doba a reklamácie", [
            "1. Záručná doba na dielo je 3 mesiace odo dňa odovzdania.",
            "2. Reklamácie sa uplatňujú písomne na adrese info@vassweb.sk.",
        ]),
        ("Článok VI – Zodpovednosť za škody", [
            "1. Zhotoviteľ nezodpovedá za škody spôsobené nesprávnym používaním diela Objednávateľom.",
            "2. Maximálna náhrada škody je obmedzená výškou zaplatenej ceny za dielo.",
        ]),
        ("Článok VII – Riešenie sporov", [
            "1. Zmluvné strany sa zaväzujú riešiť spory prednostne mimosúdnou cestou.",
            "2. Pre prípad súdneho sporu je príslušný súd podľa sídla Zhotoviteľa.",
        ]),
        ("Článok VIII – Záverečné ustanovenia", [
            "1. Táto Zmluva nadobúda platnosť a účinnosť dňom jej podpisu oboma zmluvnými stranami.",
            "2. Zmluvu možno meniť len písomnými dodatkami podpísanými oboma stranami.",
            "3. Ak niektoré ustanovenie Zmluvy je neplatné, ostatné ustanovenia zostávajú v platnosti.",
        ]),
    ]

    for heading, paragraphs in articles:
        add_gold_heading(doc, heading, level=2)
        for text in paragraphs:
            add_paragraph(doc, text, font_size=10, color=BODY_TEXT, space_after=4)

    add_signature_block(doc, left_role="konateľ, Vass & Co. s.r.o.", right_role="Objednávateľ")
    doc.save(f"{DOCS_DIR}/Vassweb-zmluva-o-dielo-sablona.docx")
    print("✓ Zmluva o dielo")


def create_hlavickovy():
    """Generate branded letterhead template."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "V Bratislave, dňa [DD.MM.RRRR]", font_size=10, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=8, space_after=24)

    add_paragraph(doc, "[Meno a priezvisko / Spoločnosť]", font_size=10, bold=True, color=BODY_TEXT, space_after=2)
    add_paragraph(doc, "[Ulica a číslo]", font_size=10, color=BODY_TEXT, space_after=2)
    add_paragraph(doc, "[PSČ Mesto]", font_size=10, color=BODY_TEXT, space_after=20)

    add_paragraph(doc, "Vec: [Predmet listu]", font_size=11, bold=True, color=BLACK, space_after=16)
    add_paragraph(doc, "Vážený pán / Vážená pani,", font_size=10, color=BODY_TEXT, space_after=12)
    add_paragraph(doc, "[Text listu...]", font_size=10, color=BODY_TEXT, space_after=24)
    add_paragraph(doc, "S úctou,", font_size=10, color=BODY_TEXT, space_after=32)

    add_paragraph(doc, "______________________________", font_size=10, color=LIGHT_GRAY, space_after=2)
    add_paragraph(doc, "Richard Vass", font_size=10, bold=True, color=BODY_TEXT, space_after=2)
    add_paragraph(doc, "konateľ, Vass & Co. s.r.o.", font_size=9, color=GRAY)

    doc.save(f"{DOCS_DIR}/Vassweb-Hlavickovy-papier.docx")
    print("✓ Hlavičkový papier")


def create_protokol():
    """Generate branded handover protocol."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "ODOVZDÁVACÍ PROTOKOL", font_size=22, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, "Protokol o odovzdaní a prevzatí diela", font_size=10, color=GRAY, space_after=12)

    # Protocol details
    make_branded_table(doc,
        headers=["Číslo protokolu", "Dátum"],
        rows=[["[ROK]-[ČÍSLO]", "[DD.MM.RRRR]"]],
    )
    add_paragraph(doc, "", space_after=8)
    add_parties_table(doc)

    add_paragraph(doc, "", space_after=8)
    add_gold_heading(doc, "Predmet odovzdania", level=2)
    make_branded_table(doc,
        headers=["", ""],
        rows=[
            ["Projekt:", "[Názov projektu]"],
            ["Popis diela:", "[Stručný popis]"],
            ["URL:", "[https://www.nazov-stranky.sk]"],
            ["Doménové meno:", "[nazov-stranky.sk]"],
            ["Hosting:", "[Poskytovateľ hostingu]"],
            ["Technológia:", "[Next.js / WordPress / iné]"],
        ],
    )

    add_paragraph(doc, "", space_after=8)
    add_gold_heading(doc, "Zoznam odovzdaných položiek", level=2)

    items = [
        "Zdrojové súbory projektu",
        "Databázové zálohy",
        "Prístupové údaje k hostingu / serveru",
        "Prístupové údaje k CMS / admin panelu",
        "Prístupové údaje k doméne / DNS",
        "Grafické podklady a zdrojové súbory dizajnu",
        "Dokumentácia projektu",
        "Návod na správu obsahu",
        "SSL certifikát",
        "Zálohovanie nastavené",
    ]
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run("☐  ")
        run.font.name = FONT_NAME; run.font.size = Pt(10); run.font.color.rgb = GOLD
        run = p.add_run(item)
        run.font.name = FONT_NAME; run.font.size = Pt(10); run.font.color.rgb = BODY_TEXT
        p.paragraph_format.space_after = Pt(3)

    add_paragraph(doc, "", space_after=8)
    add_gold_heading(doc, "Poznámky", level=2)
    for _ in range(3):
        add_paragraph(doc, "________________________________________________________________________________", font_size=10, color=LIGHT_GRAY, space_after=4)

    add_paragraph(doc, "Zmluvné strany potvrdzujú, že dielo bolo odovzdané a prevzaté v stave zodpovedajúcom dohodnutým požiadavkám.", font_size=9.5, color=BODY_TEXT, space_before=12)

    add_signature_block(doc, left_role="konateľ, Vass & Co. s.r.o.", right_role="Objednávateľ")
    doc.save(f"{DOCS_DIR}/Vassweb-Odovzdaci-protokol.docx")
    print("✓ Odovzdávací protokol")


def create_nda():
    """Generate branded NDA template."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "DOHODA O MLČANLIVOSTI", font_size=22, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, "uzatvorená podľa § 269 ods. 2 Obchodného zákonníka SR", font_size=10, color=GRAY, space_after=16)

    add_parties_table(doc)
    add_paragraph(doc, "sa dohodli na uzatvorení tejto Dohody o mlčanlivosti (ďalej len \u201eDohoda\u201c):", font_size=10, color=BODY_TEXT, space_before=12, space_after=12)

    articles = [
        ("Článok I – Predmet Dohody", [
            "1. Táto Dohoda upravuje podmienky zachovania mlčanlivosti o dôverných informáciách, ktoré si zmluvné strany poskytnú alebo sprístupnia v súvislosti so vzájomnou spoluprácou.",
            "2. Dôvernými informáciami sa rozumejú akékoľvek informácie, dokumenty, technické, obchodné, finančné alebo iné údaje, vrátane know-how, zdrojových kódov, databáz, cenových politík a interných procesov.",
        ]),
        ("Článok II – Záväzky zmluvných strán", None),
        ("Článok III – Výnimky z mlčanlivosti", None),
        ("Článok IV – Trvanie Dohody", [
            "1. Táto Dohoda je uzatvorená na dobu neurčitú.",
            "2. Záväzky mlčanlivosti trvajú aj po ukončení obchodnej spolupráce zmluvných strán, a to po dobu 5 (päť) rokov od ukončenia spolupráce.",
            "3. Každá zo zmluvných strán môže Dohodu písomne vypovedať s výpovednou lehotou 30 dní.",
        ]),
        ("Článok V – Sankcie za porušenie", [
            "1. V prípade porušenia záväzku mlčanlivosti je povinná zmluvná strana nahradiť druhej zmluvnej strane preukázanú škodu v plnom rozsahu.",
            "2. Zmluvné strany sa dohodli na zmluvnej pokute vo výške 5 000 € za každý jednotlivý prípad porušenia tejto Dohody, čím nie je dotknutý nárok na náhradu škody.",
        ]),
        ("Článok VI – Záverečné ustanovenia", [
            "1. Táto Dohoda nadobúda platnosť a účinnosť dňom jej podpisu oboma zmluvnými stranami.",
            "2. Dohodu možno meniť len písomnými dodatkami podpísanými oboma stranami.",
            "3. Táto Dohoda sa riadi právnym poriadkom Slovenskej republiky.",
        ]),
    ]

    for heading, paragraphs in articles:
        add_gold_heading(doc, heading, level=2)
        if heading == "Článok II – Záväzky zmluvných strán":
            add_paragraph(doc, "1. Každá zmluvná strana sa zaväzuje:", font_size=10, color=BODY_TEXT, space_after=4)
            add_bullet_list(doc, [
                "uchovávať dôverné informácie v tajnosti a neposkytovať ich tretím osobám bez predchádzajúceho písomného súhlasu druhej zmluvnej strany,",
                "používať dôverné informácie výlučne na účely spolupráce podľa tejto Dohody,",
                "zabezpečiť primeranú ochranu dôverných informácií pred neoprávneným prístupom alebo zverejnením,",
                "informovať druhú zmluvnú stranu o akomkoľvek neoprávnenom sprístupnení alebo úniku dôverných informácií.",
            ])
        elif heading == "Článok III – Výnimky z mlčanlivosti":
            add_paragraph(doc, "1. Záväzok mlčanlivosti sa nevzťahuje na informácie, ktoré:", font_size=10, color=BODY_TEXT, space_after=4)
            add_bullet_list(doc, [
                "boli verejne dostupné pred podpisom tejto Dohody alebo sa stali verejne dostupnými bez porušenia tejto Dohody,",
                "boli zákonite v držbe príjemcu ešte pred ich poskytnutím poskytovateľom,",
                "boli nezávisle vyvinuté príjemcom bez použitia dôverných informácií,",
                "museli byť zverejnené na základe zákonnej povinnosti alebo rozhodnutia príslušného orgánu.",
            ])
        elif paragraphs:
            for text in paragraphs:
                add_paragraph(doc, text, font_size=10, color=BODY_TEXT, space_after=4)

    add_signature_block(doc, left_role="konateľ, Vass & Co. s.r.o.", right_name="[Meno a priezvisko]", right_role="Prijímateľ")
    doc.save(f"{DOCS_DIR}/Vassweb-NDA-Dohoda-o-mlcanlivosti.docx")
    print("✓ NDA")


def create_vop():
    """Generate branded Terms of Service."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "VŠEOBECNÉ OBCHODNÉ PODMIENKY", font_size=20, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, "Vass & Co. s.r.o. — účinné od 01.01.2025", font_size=10, color=GRAY, space_after=16)

    articles = {
        "Článok 1 – Úvodné ustanovenia": [
            f"1.1 Tieto Všeobecné obchodné podmienky (ďalej \u201eVOP\u201c) upravujú zmluvné vzťahy medzi spoločnosťou VVD s. r. o. (VASSWEB) so sídlom v Bratislave, IČO: {COMPANY_INFO['ico']} (ďalej \u201ePoskytovate\u013e\u201c) a jeho klientmi (ďalej \u201eKlient\u201c).",
            "1.2 VOP sú platné pre všetky zmluvy, objednávky a cenové ponuky uzatvorené medzi Poskytovateľom a Klientom, pokiaľ sa strany písomne nedohodnú inak.",
            "1.3 Odoslaním záväznej objednávky alebo podpisom zmluvy Klient vyjadruje súhlas s týmito VOP.",
        ],
        "Článok 2 – Predmet služieb": None,
        "Článok 3 – Objednávka a uzatvorenie zmluvy": [
            "3.1 Zmluva sa uzatvára na základe písomnej cenovej ponuky vystavenej Poskytovateľom a jej záväzného potvrdenia Klientom.",
            "3.2 Cenová ponuka je platná 30 dní od dátumu vystavenia, ak nie je uvedené inak.",
            "3.3 Zmeny v rozsahu diela po podpise zmluvy sú možné len na základe písomného dodatku k zmluve.",
        ],
        "Článok 4 – Ceny a platobné podmienky": [
            "4.1 Ceny sú stanovené dohodou v cenovej ponuke alebo zmluve. Poskytovateľ je platcom DPH.",
            "4.2 Štandardné platobné podmienky: záloha 50 % pred zahájením prác, doplatenie 50 % po odovzdaní diela, splatnosť faktúr 14 dní od vystavenia.",
            "4.3 Pri omeškaní platby si Poskytovateľ vyhradzuje právo účtovať zmluvný úrok z omeškania vo výške 0,05 % z dlžnej sumy za každý deň omeškania.",
            f"4.4 Platby sa realizujú bezhotovostným prevodom na IBAN: {COMPANY_INFO['iban']}.",
        ],
        "Článok 5 – Dodacie podmienky a termíny": [
            "5.1 Termín dodania je dohodnutý v cenovej ponuke alebo zmluve a začína plynúť od uhradenia zálohy a dodania všetkých potrebných podkladov.",
            "5.2 V prípade omeškania Klienta s dodaním podkladov alebo s platbou sa termín primerane predlžuje.",
        ],
        "Článok 6 – Odovzdanie diela a reklamácie": [
            "6.1 Dielo sa odovzdáva elektronicky alebo podpisom odovzdacieho protokolu.",
            "6.2 Klient je povinný skontrolovať dielo do 5 pracovných dní od odovzdania.",
            "6.3 Záručná doba na dielo je 3 mesiace od odovzdania.",
            "6.4 Reklamácie sa uplatňujú písomne na adrese info@vassweb.sk.",
        ],
        "Článok 7 – Autorské práva a licencia": [
            "7.1 Po úplnom uhradení ceny diela prechádzajú na Klienta majetkové autorské práva k dielu.",
            "7.2 Poskytovateľ si vyhradzuje právo uviesť realizované dielo vo svojom portfóliu.",
        ],
        "Článok 8 – Ochrana osobných údajov": [
            "8.1 Poskytovateľ spracúva osobné údaje v súlade s nariadením GDPR (EÚ) 2016/679 a zákonom č. 18/2018 Z. z.",
            "8.2 Osobné údaje sú spracúvané výlučne za účelom plnenia zmluvy a nie sú poskytované tretím stranám bez súhlasu.",
        ],
        "Článok 9 – Záverečné ustanovenia": [
            "9.1 Tieto VOP nadobúdajú účinnosť dňom ich zverejnenia na webovej stránke vassweb.sk.",
            "9.2 Poskytovateľ si vyhradzuje právo VOP zmeniť, pričom o zmene informuje Klienta najmenej 14 dní vopred.",
            "9.3 Na právne vzťahy neupravené týmito VOP sa vzťahujú príslušné ustanovenia Obchodného zákonníka SR.",
        ],
    }

    services = [
        "tvorba webových stránok a e-shopov",
        "webový dizajn a UX/UI návrhy",
        "SEO optimalizácia a digitálny marketing",
        "správa a údržba webových projektov",
        "hosting a technická podpora",
        "vývoj webových aplikácií na mieru",
        "AI riešenia a automatizácia procesov",
    ]

    for heading, paragraphs in articles.items():
        add_gold_heading(doc, heading, level=2)
        if heading == "Článok 2 – Predmet služieb":
            add_paragraph(doc, "2.1 Poskytovateľ poskytuje najmä tieto služby:", font_size=10, color=BODY_TEXT, space_after=4)
            add_bullet_list(doc, services)
        elif paragraphs:
            for text in paragraphs:
                add_paragraph(doc, text, font_size=10, color=BODY_TEXT, space_after=4)

    doc.save(f"{DOCS_DIR}/Vassweb-Vseobecne-obchodne-podmienky.docx")
    print("✓ VOP")


def create_ponukovy_list(title, subtitle, intro_text, features, tables_data, contact_text, output_path):
    """Generic function for creating branded ponukový list."""
    doc = Document()
    add_header_block(doc)
    add_footer_block(doc)

    add_paragraph(doc, "PONUKOVÝ LIST", font_size=11, bold=True, color=GOLD_DARK, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, title, font_size=20, bold=True, color=BLACK, space_after=2)
    add_paragraph(doc, subtitle, font_size=10, color=GRAY, space_after=16)

    if intro_text:
        add_paragraph(doc, intro_text, font_size=10, color=BODY_TEXT, space_after=12)

    if features:
        for section_title, items in features:
            add_gold_heading(doc, section_title, level=2)
            add_bullet_list(doc, items)
            add_paragraph(doc, "", space_after=4)

    for table_title, headers, rows in tables_data:
        if table_title:
            add_gold_heading(doc, table_title, level=2)
        make_branded_table(doc, headers, rows)
        add_paragraph(doc, "", space_after=8)

    # Disclaimer
    add_paragraph(doc, "Uvedené ceny sú orientačné a slúžia ako rámcový prehľad. Každú zákazku posudzujeme individuálne — presnú cenovú ponuku vám pripravíme do 24 hodín po obdržaní podkladov.", font_size=8.5, color=GRAY, space_before=12, space_after=12)

    if contact_text:
        add_gold_heading(doc, "Kontakt a ďalší postup", level=2)
        add_paragraph(doc, contact_text, font_size=10, color=BODY_TEXT, space_after=4)

    add_paragraph(doc, f"Richard Vass  |  {COMPANY_INFO['phone']}  |  {COMPANY_INFO['email']}  |  www.vassco.sk", font_size=9, color=GRAY, space_after=2)
    add_paragraph(doc, f"VVD s.r.o. (Vass & Co.)  |  IČO: {COMPANY_INFO['ico']}  |  DIČ: {COMPANY_INFO['dic']}  |  IČ DPH: {COMPANY_INFO['ic_dph']}", font_size=8, color=GRAY)

    doc.save(output_path)
    print(f"✓ {os.path.basename(output_path)}")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("═══ VASSWEB REBRAND — Generating documents ═══\n")

    # 1. Vassweb documents
    create_faktura()
    create_cenova_ponuka()
    create_zmluva()
    create_hlavickovy()
    create_protokol()
    create_nda()
    create_vop()

    # 2. Ponukové listy
    PL_DIR = f"{DRIVE}/02_Ponukové listy"

    create_ponukovy_list(
        title="Facility Management",
        subtitle="Pravidelné upratovanie kancelárií  |  B2B Paušál  |  2026",
        intro_text=None,
        features=[
            ("Prečo Vass & Co. Facility?", [
                "Stabilné pravidelné upratovanie kancelárií v Bratislave a okolí",
                "Flexibilná frekvencia — 1× až 5× týždenne podľa vašich potrieb",
                "Fotodokumentácia a mesačný report kvality na požiadanie",
                "Hygienický servis (papier, mydlo, spotrebák) v cene doplnku",
                "Jeden kontakt, jedna zmluva, žiadne prekvapenia",
            ]),
            ("Čo je súčasťou každej návštevy?", [
                "Vysávanie kobercov a umývanie podláh",
                "Utieranie prachu zo všetkých dostupných plôch",
                "Sanitácia WC a sociálnych zariadení",
                "Čistenie sklenených plôch a dverí",
                "Vynesenie odpadkov a výmena vriec",
                "Dezinfekcia dotykových plôch (kľučky, svetlá, výťahové tlačidlá)",
            ]),
        ],
        tables_data=[
            ("Cenník — € / m² / mesiac (bez DPH)", ["Plocha objektu", "1× týžd.", "2× týžd.", "3× týžd.", "5× týžd."], [
                ["do 150 m²", "od 1,05 €", "od 1,25 €", "od 1,50 €", "od 2,10 €"],
                ["150 – 500 m²", "od 0,85 €", "od 1,10 €", "od 1,35 €", "od 1,85 €"],
                ["500 – 1 500 m²", "od 0,70 €", "od 0,95 €", "od 1,20 €", "od 1,55 €"],
            ]),
            ("Doplnky (bez DPH)", ["Doplnok / služba", "Cena", "Jednotka"], [
                ["WC / toaleta", "od 15 €", "€ / WC / mesiac"],
                ["Kuchynka malá", "od 25 €", "€ / mesiac"],
                ["Kuchynka veľká", "od 50 €", "€ / mesiac"],
                ["Hygienický servis", "od 30 €", "€ / mesiac"],
                ["Mesačný report kvality", "od 20 €", "€ / mesiac"],
            ]),
        ],
        contact_text="Kontaktujte nás — nezáväznú cenovú ponuku pripravíme do 24 hodín.",
        output_path=f"{PL_DIR}/PL_Facility_B2B.docx",
    )

    create_ponukovy_list(
        title="Tepovanie",
        subtitle="Extrakčné čistenie kobercov, sedačiek a matracov  |  B2B + B2C  |  2026",
        intro_text=None,
        features=[
            ("Prečo profesionálne tepovanie?", [
                "Hlboké extrakčné čistenie odstraňuje nečistoty, ktoré bežný vysávač nedostane",
                "Predlžuje životnosť kobercov, sedačiek a matracov",
                "Eliminuje alergény, roztoče a baktérie",
                "Fotodokumentácia pred aj po — vidíte výsledok",
                "Test chémie na malej ploche pred každou zákazkou",
                "Pokyny na schnutie a starostlivosť po čistení",
            ]),
            ("Postup pri každej zákazke", [
                "Obhliadka a identifikácia materiálu a škvrn",
                "Fotodokumentácia stavu PRED čistením",
                "Test chémie na skrytom mieste (bezpečnosť farby)",
                "Extrakčné čistenie zvolenou metódou a chémiou",
                "Fotodokumentácia stavu PO čistení",
                "Odovzdanie protokolu + pokyny na schnutie",
            ]),
        ],
        tables_data=[
            ("Cenník tepovanie (bez DPH)", ["Predmet čistenia", "Cena bez DPH", "Jednotka"], [
                ["Koberce", "od 1,80 €", "€ / m²"],
                ["Kancelárska stolička (sedák + operadlo)", "od 10 €", "€ / ks"],
                ["Kancelárske kreslo", "od 18 €", "€ / ks"],
                ["Sedačka — 2-miestna", "od 35 €", "€ / ks"],
                ["Sedačka — 3-miestna", "od 50 €", "€ / ks"],
                ["Sedačka — rohová", "od 70 €", "€ / ks"],
                ["Matrac — 1-lôžko", "od 25 €", "€ / ks"],
                ["Matrac — 2-lôžko", "od 40 €", "€ / ks"],
                ["Autosedačka", "od 15 €", "€ / ks"],
            ]),
        ],
        contact_text="Minimálna objednávka: 100 € bez DPH. K cene sa pripočíta DPH 23 %.",
        output_path=f"{PL_DIR}/PL_Tepovanie.docx",
    )

    create_ponukovy_list(
        title="Airbnb Turnover",
        subtitle="Profesionálna príprava bytu medzi hosťami  |  B2B + B2C  |  2026",
        intro_text="Kompletná príprava vášho bytu medzi odchodom a príchodom hostí — upratovanie, výmena bielizne, doplnenie spotrebáku a fotoreport.",
        features=[
            ("Čo zahŕňa turnover?", [
                "Rýchla a spoľahlivá príprava bytu medzi hosťami",
                "Fotoreport (5–10 fotiek) po každom turnoveri — dôkaz kvality",
                "Výmena bielizne a uterákov (linen exchange)",
                "Doplnenie spotrebáku z vašich zásob",
                "Expres servis do 6 hodín od požiadavky",
                "Koordinácia s vaším Airbnb kalendárom",
            ]),
        ],
        tables_data=[
            ("Cenník Turnover (bez DPH)", ["Služba", "Jednotka", "Cena bez DPH"], [
                ["Turnover STANDARD — 1-izb byt", "€ / turnover", "od 45 €"],
                ["Turnover STANDARD — 2-izb byt", "€ / turnover", "od 65 €"],
                ["Turnover STANDARD — 3-izb byt", "€ / turnover", "od 85 €"],
                ["Turnover PREMIUM — 1-izb byt", "€ / turnover", "od 59 €"],
                ["Turnover PREMIUM — 2-izb byt", "€ / turnover", "od 82 €"],
                ["Turnover PREMIUM — 3-izb byt", "€ / turnover", "od 109 €"],
            ]),
            ("Bielizeň a doplnky (bez DPH)", ["Položka", "Jednotka", "Cena bez DPH"], [
                ["Výmena bielizne — hostiteľove zásoby", "€ / lôžko", "od 8 €"],
                ["Uteráky set — hostiteľove zásoby", "€ / set", "od 5 €"],
                ["Linen exchange (naša bielizeň)", "€ / lôžko", "od 22 €"],
                ["Uteráky (naše zásoby)", "€ / set", "od 12 €"],
                ["Welcome kit (káva, čaj, cukríky)", "€ / set", "od 8 €"],
            ]),
            ("Príplatky", ["Príplatok", "Sadzba", "Podmienka"], [
                ["Expres servis do 6 hodín", "+25 %", "prirážka k balíku"],
                ["Nedeľa / sviatok", "+20 %", "prirážka k balíku"],
                ["Nočný turnover (22:00–06:00)", "+30 %", "prirážka k balíku"],
            ]),
        ],
        contact_text="Spravujete jeden alebo viacero bytov? Ozvite sa — nastavíme spoluprácu presne podľa vašich potrieb.",
        output_path=f"{PL_DIR}/PL_Airbnb_Turnover.docx",
    )

    create_ponukovy_list(
        title="Špeciálne služby",
        subtitle="Upratovanie po rekonštrukcii, sťahovaní a ďalšie  |  B2B + B2C  |  2026",
        intro_text="Okrem pravidelných služieb ponúkame aj jednorázové špeciálne upratovanie pre situácie, kedy bežné riešenie nestačí.",
        features=[
            ("Jednorázové a špeciálne upratovanie", [
                "Upratovanie po rekonštrukcii a stavebných prácach",
                "Upratovanie pred/po sťahovaní",
                "Hĺbkové čistenie pred kolaudáciou",
                "Čistenie po poistných udalostiach (vytopenie, požiar)",
                "Dezinfekcia a sanitácia priestorov",
                "Čistenie po evente / konferencii",
            ]),
            ("Ako to funguje?", [
                "Zašlite nám fotky priestoru alebo si dohodneme obhliadku",
                "Do 24 hodín dostanete cenovú ponuku",
                "Po odsúhlasení nasadíme tím v dohodnutom termíne",
                "Fotodokumentácia pred a po čistení",
                "Fakturácia po dokončení s detailným popisom prác",
            ]),
        ],
        tables_data=[
            ("Cenník špeciálnych služieb (bez DPH)", ["Služba", "Cena", "Jednotka"], [
                ["Upratovanie po rekonštrukcii — BASIC", "od 3,50 €", "€ / m²"],
                ["Upratovanie po rekonštrukcii — DEEP", "od 5,50 €", "€ / m²"],
                ["Upratovanie pred/po sťahovaní", "od 3,00 €", "€ / m²"],
                ["Hĺbkové čistenie pred kolaudáciou", "od 4,50 €", "€ / m²"],
                ["Dezinfekcia priestorov", "od 2,00 €", "€ / m²"],
                ["Čistenie po evente", "od 3,00 €", "€ / m²"],
                ["Umývanie okien (jednorazové)", "od 4,00 €", "€ / okno"],
                ["Čistenie fasády / exteriéru", "od 5,00 €", "€ / m²"],
                ["Parovanie a dezinfekcia", "od 3,50 €", "€ / m²"],
            ]),
        ],
        contact_text="Minimálna objednávka: 150 € bez DPH. K cene sa pripočíta DPH 23 %.",
        output_path=f"{PL_DIR}/PL_Specialne_Sluzby.docx",
    )

    create_ponukovy_list(
        title="Správa bytových domov",
        subtitle="Kompletná údržba a čistenie spoločných priestorov  |  B2B  |  2026",
        intro_text="Ponúkame pravidelné aj jednorazové upratovanie spoločných priestorov bytových domov.",
        features=[
            ("Komplexná starostlivosť o váš bytový dom", [
                "Pravidelné upratovanie schodísk, chodieb a vstupných priestorov",
                "Umývanie okien, zábradlí a výťahov",
                "Dezinfekcia dotykových plôch a kľučiek",
                "Zimná údržba — posyp, odhŕňanie snehu",
                "Mesačný report s fotodokumentáciou",
                "Zastupiteľnosť — žiadne výpadky v upratovaní",
            ]),
            ("Prečo Vass & Co. pre váš bytový dom?", [
                "Stabilný tím — nestriedame pracovníkov, poznáme váš dom",
                "Vlastné vybavenie a profesionálna čistiaca chémia",
                "Fakturácia mesačne na správcu alebo SVB",
                "Rámcová zmluva s jasne definovaným rozsahom",
                "Flexibilný prístup — reagujeme na vaše požiadavky do 24h",
            ]),
        ],
        tables_data=[
            ("Cenník — pravidelné upratovanie (bez DPH)", ["Služba", "Frekvencia", "Cena", "Jednotka"], [
                ["ŠTANDARD — spoločné priestory", "1× týždenne", "od 6 €", "€ / byt / mesiac"],
                ["ŠTANDARD — spoločné priestory", "2× týždenne", "od 8 €", "€ / byt / mesiac"],
                ["PRÉMIUM — vrátane okien a výťahov", "1× týždenne", "od 9 €", "€ / byt / mesiac"],
                ["PRÉMIUM — vrátane okien a výťahov", "2× týždenne", "od 12 €", "€ / byt / mesiac"],
                ["KOMPLET — vrátane zimnej údržby", "dohodou", "individuálne", "€ / byt / mesiac"],
            ]),
            ("Doplnkové služby", ["Služba", "Frekvencia", "Cena"], [
                ["Umývanie okien (spoločné priestory)", "1× štvrťročne", "od 3 €/okno"],
                ["Hĺbkové čistenie podláh", "1× polročne", "od 1,50 €/m²"],
                ["Dezinfekcia spoločných priestorov", "1× mesačne", "od 0,80 €/m²"],
                ["Zimný posyp + odhŕňanie", "sezóna", "od 2 €/byt/mesiac"],
                ["Čistenie garážových priestorov", "1× štvrťročne", "individuálne"],
            ]),
        ],
        contact_text="Minimum 200 € / mesiac bez DPH. K cene sa pripočíta DPH 23 %.",
        output_path=f"{PL_DIR}/PL_Sprava_Bytovych_Domov.docx",
    )

    create_ponukovy_list(
        title="Kompletný prehľad služieb Vass & Co.",
        subtitle="Facility | Tepovanie | Airbnb Turnover  |  B2B + B2C  |  2026",
        intro_text="Vass & Co. (VVD s.r.o.) je slovenská upratovacia spoločnosť so sídlom v Bratislave. Poskytujeme B2B aj B2C klientom kompletné upratovacie služby.",
        features=[
            ("Kto sme", [
                "Profesionálny tím s vlastným vybavením a chémiou",
                "Fotodokumentácia a reporty ku každej zákazke",
                "Flexibilita — jednorazovo aj dlhodobá zmluva",
                "Jeden kontakt pre všetky vaše priestory",
            ]),
        ],
        tables_data=[
            ("1 — Facility: Pravidelné upratovanie kancelárií (B2B)", ["Plocha / Frekvencia", "1× / týžd.", "2× / týžd.", "3× / týžd.", "5× / týžd."], [
                ["do 150 m² (€/m²/mes.)", "od 1,05 €", "od 1,25 €", "od 1,50 €", "od 2,10 €"],
                ["150–500 m²", "od 0,85 €", "od 1,10 €", "od 1,35 €", "od 1,85 €"],
                ["500–1 500 m²", "od 0,70 €", "od 0,95 €", "od 1,20 €", "od 1,55 €"],
            ]),
            ("2 — Tepovanie: Extrakčné čistenie (B2B + B2C)", ["Predmet", "Cena bez DPH", "Jednotka"], [
                ["Koberce", "od 1,80 €", "€ / m²"],
                ["Kancelárska stolička", "od 10 €", "€ / ks"],
                ["Sedačka 2-miestna", "od 35 €", "€ / ks"],
                ["Sedačka rohová", "od 70 €", "€ / ks"],
                ["Matrac 2-lôžko", "od 40 €", "€ / ks"],
            ]),
            ("3 — Airbnb Turnover: Príprava bytu medzi hosťami", ["Balík", "1-izb", "2-izb", "3-izb"], [
                ["Turnover STANDARD (bez DPH)", "od 45 €", "od 65 €", "od 85 €"],
                ["Turnover PREMIUM + fotoreport", "od 59 €", "od 82 €", "od 109 €"],
            ]),
        ],
        contact_text="Nezáväzná cenová ponuka do 24 hodín po obdržaní podkladov (plocha, frekvencia, adresa).",
        output_path=f"{PL_DIR}/PL_Kompletny_Prehlad.docx",
    )

    # 3. Cenník 2026
    create_ponukovy_list(
        title="VASS & CO.",
        subtitle="Facility  ·  Tepovanie  ·  Airbnb Turnover  |  Kompletný cenník — platnosť od 07.03.2026",
        intro_text="Tento dokument obsahuje finálny návrh cien pre všetky tri segmenty Vass & Co. Ceny sú nastavené na vyváženo–prémiovú pozíciu na bratislavskom trhu.",
        features=[],
        tables_data=[
            ("01  Facility Management — € / m² / mesiac (bez DPH)", ["Plocha / Frekvencia", "1× / týžd.", "2× / týžd.", "3× / týžd.", "5× / týžd."], [
                ["do 150 m²", "1,05 €", "1,30 €", "1,60 €", "2,25 €"],
                ["150 – 500 m²", "0,90 €", "1,15 €", "1,45 €", "1,95 €"],
                ["500 – 1 500 m²", "0,75 €", "1,00 €", "1,30 €", "1,65 €"],
                ["nad 1 500 m²", "individuálne", "individuálne", "individuálne", "individuálne"],
            ]),
            ("Doplnky k paušálu", ["Doplnok / služba", "Cena bez DPH", "Jednotka"], [
                ["WC / toaleta", "18 €", "€ / WC / mesiac"],
                ["Kuchynka malá", "30 €", "€ / mesiac"],
                ["Kuchynka veľká", "55 €", "€ / mesiac"],
                ["Hygienický servis", "35 €", "€ / mesiac"],
                ["Mesačný report kvality", "25 €", "€ / mesiac"],
            ]),
            ("02  Tepovanie (bez DPH)", ["Predmet čistenia", "Cena bez DPH", "Jednotka"], [
                ["Koberec (extrakčné čistenie)", "1,85 €", "€ / m²"],
                ["Behúň / kusový koberec", "25 €", "€ / ks"],
                ["Kancelárska stolička", "12 €", "€ / ks"],
                ["Kancelárske kreslo", "20 €", "€ / ks"],
                ["Sedačka — 2-miestna", "39 €", "€ / ks"],
                ["Sedačka — 3-miestna", "55 €", "€ / ks"],
                ["Sedačka — rohová L/U", "75 €", "€ / ks"],
                ["Matrac — 1-lôžko", "29 €", "€ / ks"],
                ["Matrac — 2-lôžko", "45 €", "€ / ks"],
                ["Autosedačka (1 ks)", "18 €", "€ / ks"],
                ["Komplet auto interiér", "89 €", "€ / auto"],
            ]),
            ("03  Airbnb Turnover (bez DPH)", ["Služba", "1-izb byt", "2-izb byt", "3-izb byt"], [
                ["Turnover STANDARD", "49 €", "69 €", "89 €"],
                ["Turnover PREMIUM (vrátane fotoreportu)", "64 €", "89 €", "114 €"],
                ["Linen exchange — naša bielizeň", "22 € / lôžko", "22 € / lôžko", "22 € / lôžko"],
                ["Welcome kit (káva, čaj, cukríky)", "9 €", "9 €", "9 €"],
                ["Expres servis do 6h", "+25 %", "+25 %", "+25 %"],
                ["Nedeľa / sviatok", "+20 %", "+20 %", "+20 %"],
            ]),
        ],
        contact_text=None,
        output_path=f"{DRIVE}/03_Dokumenty/VassCo Novy Cennik 2026.docx",
    )

    print("\n═══ DONE — All documents rebranded! ═══")
